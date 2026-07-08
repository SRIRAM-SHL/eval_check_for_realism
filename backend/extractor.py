"""
extractor.py - Pull clean plain text out of uploaded documents.

Supports .docx, .pdf, .txt and .md. Everything is normalised into a single
string that the ingest stage can chunk. (Note: V2's extractor was missing its
PDF reader; this one implements `_from_pdf` correctly.)
"""
from __future__ import annotations

import io
import re

from docx import Document
from pypdf import PdfReader


class UnsupportedFileError(ValueError):
    """Raised when a file extension we don't handle is uploaded."""


def _clean(text: str) -> str:
    """Collapse noisy whitespace while keeping paragraph breaks."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            # Keep heading styles as markdown so ingest can detect sections.
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                level = "".join(c for c in style if c.isdigit()) or "1"
                parts.append(f"{'#' * min(int(level), 6)} {para.text.strip()}")
            else:
                parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - skip unreadable pages, keep the rest
            txt = ""
        if txt.strip():
            parts.append(txt)
    return "\n\n".join(parts)


def _from_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on file extension and return cleaned plain text."""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        raw = _from_docx(data)
    elif name.endswith(".pdf"):
        raw = _from_pdf(data)
    elif name.endswith((".txt", ".md")):
        raw = _from_txt(data)
    elif name.endswith(".doc"):
        raise UnsupportedFileError(
            "Legacy .doc files are not supported. Please save as .docx and re-upload."
        )
    else:
        raise UnsupportedFileError(
            f"Unsupported file type: {filename!r}. Upload a .docx, .pdf, .txt or .md file."
        )

    cleaned = _clean(raw)
    if not cleaned:
        raise UnsupportedFileError(
            "No readable text was found in the document (it may be scanned images)."
        )
    return cleaned
